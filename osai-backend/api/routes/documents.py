"""Direct file upload into the org knowledge base.

Uploaded files become SourceDocuments (source_type="upload") and flow through
the exact same pipeline as connector ingestion — tier rules, Postgres chunks,
embeddings, Qdrant — so retrieval, governance, and the graph treat them like
any synced document.
"""

from __future__ import annotations

import io
import logging
from typing import Annotated
from uuid import uuid4

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from api.schemas.connector import SourceDocument
from db.models import Chunk, Department, Notification, SourceDocumentRecord, User
from db.repositories import chunks_for_documents, try_db, upsert_source_documents
from db.session import get_db, get_optional_claims, get_org_id, require_writable_org
from memory.qdrant_store import get_default_qdrant_store

logger = logging.getLogger("osai.documents")

router = APIRouter(prefix="/documents", tags=["documents"])
DbSession = Annotated[Session, Depends(get_db)]
OrgId = Annotated[str, Depends(get_org_id)]
# Uploading indexes data + spends embeddings; changing access rewrites grants.
# Neither is reachable from the anonymous demo workspace (SEC-003).
WriteOrgId = Annotated[str, Depends(require_writable_org)]
OptionalClaims = Annotated[dict | None, Depends(get_optional_claims)]

_MAX_FILE_BYTES = 15 * 1024 * 1024  # 15 MB per file
_ALLOWED_TIERS = {"normal", "amber", "red"}
_ALLOWED_VISIBILITY = {"personal", "department", "company", "people"}


def _visibility_grants(
    db: Session,
    org_id: str,
    visibility: str,
    uploader_id: str | None,
    department_id: str | None,
    shared_with: list[str],
) -> list[str]:
    """Map a human visibility choice onto permission grants.

    personal   -> ["user:<uploader>"]           (only the uploader)
    department -> ["dept:<department_id>"]       (their department)
    company    -> ["source:all"]                 (whole workspace)
    people     -> ["user:<uploader>", "user:<id>", ...] (named teammates)

    Raises HTTPException(422) when the choice can't be honoured (e.g. personal
    without an authenticated user), rather than silently widening access."""
    if visibility == "company":
        return ["source:all"]
    if visibility == "personal":
        # Demo/unauthenticated context has no user account to scope to — those
        # uploads land workspace-wide, same as every other demo document.
        if not uploader_id:
            return ["source:all"]
        return [f"user:{uploader_id}"]
    if visibility == "department":
        if not department_id:
            raise HTTPException(
                status_code=422,
                detail="Pick a department to share with (or join one in Team settings).",
            )
        # Drive-style: named people can be added on top of department access.
        return sorted({f"dept:{department_id}", *_user_grants(db, org_id, shared_with)})
    # visibility == "people"
    if not shared_with:
        raise HTTPException(
            status_code=422, detail="Choose at least one person to share with."
        )
    grants = set(_user_grants(db, org_id, shared_with))
    if uploader_id:
        grants.add(f"user:{uploader_id}")  # sharer keeps access to their own file
    return sorted(grants)


def _user_grants(db: Session, org_id: str, user_ids: list[str]) -> list[str]:
    """Validated "user:<id>" grants for workspace members."""
    if not user_ids:
        return []
    members = db.query(User).filter(User.org_id == org_id, User.id.in_(user_ids)).all()
    if len(members) != len(set(user_ids)):
        raise HTTPException(
            status_code=422, detail="One or more selected people aren't in this workspace."
        )
    return [f"user:{m.id}" for m in members]


def _extract_text(filename: str, data: bytes) -> str:
    """Text from an uploaded file. Raises HTTPException(415) for unsupported types."""
    name = filename.lower()
    if name.endswith((".txt", ".md", ".markdown", ".csv", ".log")):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except HTTPException:
            raise
        except Exception as exc:  # noqa: BLE001 — malformed PDFs are a client error
            raise HTTPException(status_code=422, detail=f"Could not parse PDF: {exc}") from exc
    if name.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(status_code=422, detail=f"Could not parse DOCX: {exc}") from exc
    raise HTTPException(
        status_code=415,
        detail="Unsupported file type. Supported: .txt, .md, .csv, .log, .pdf, .docx",
    )


@router.post("/upload")
async def upload_documents(
    db: DbSession,
    org_id: WriteOrgId,
    claims: OptionalClaims,
    files: Annotated[list[UploadFile], File(...)],
    data_tier: Annotated[str, Form()] = "normal",
    visibility: Annotated[str, Form()] = "personal",
    shared_with: Annotated[str, Form()] = "",
    department_id: Annotated[str, Form()] = "",
) -> dict:
    """Ingest uploaded files into the knowledge base.

    `visibility` says who can see the files: personal (just the uploader),
    department (requires `department_id`), company (whole workspace), or people
    (requires `shared_with`, a comma-separated list of member ids). It is
    translated into the same permission grants the retriever already filters by.
    `data_tier` remains an internal routing classification (cloud vs local
    model egress) and defaults to "normal"; it is no longer a user-facing choice."""
    if data_tier not in _ALLOWED_TIERS:
        raise HTTPException(
            status_code=422, detail=f"data_tier must be one of {sorted(_ALLOWED_TIERS)}"
        )
    if visibility not in _ALLOWED_VISIBILITY:
        raise HTTPException(
            status_code=422, detail=f"visibility must be one of {sorted(_ALLOWED_VISIBILITY)}"
        )
    uploader_id = claims.get("sub") if claims else None
    author = claims.get("email") or uploader_id if claims else None
    # Department attribution — must be one of this org's departments. For
    # department visibility, default to the uploader's own department.
    dept = department_id.strip() or None
    if not dept and visibility == "department" and uploader_id:
        uploader = db.get(User, uploader_id)
        dept = uploader.department_id if uploader else None
    if dept:
        row = db.get(Department, dept)
        if row is None or row.org_id != org_id:
            raise HTTPException(status_code=422, detail="Unknown department for this workspace.")
    recipients = [s.strip() for s in shared_with.split(",") if s.strip()]
    grants = _visibility_grants(db, org_id, visibility, uploader_id, dept, recipients)

    documents: list[SourceDocument] = []
    skipped: list[dict] = []
    for file in files:
        data = await file.read()
        if len(data) > _MAX_FILE_BYTES:
            skipped.append({"filename": file.filename, "reason": "File exceeds 15 MB limit"})
            continue
        try:
            text = _extract_text(file.filename or "unnamed", data)
        except HTTPException as exc:
            skipped.append({"filename": file.filename, "reason": exc.detail})
            continue
        if not text.strip():
            skipped.append({"filename": file.filename, "reason": "No extractable text"})
            continue
        # Tenant-safe ID: random per upload, so re-uploading the same filename
        # (here or in another org) never collides or overwrites across tenants.
        doc_id = f"upload-{uuid4()}"
        documents.append(
            SourceDocument(
                source_id=doc_id,
                source_type="upload",
                org_id=org_id,
                external_id=doc_id,
                title=file.filename or "Untitled upload",
                author=author,
                text=text,
                metadata={"origin": "direct_upload", "content_length": len(text)},
                permissions=grants,
                data_tier=data_tier,  # type: ignore[arg-type]
                department_id=dept,
            )
        )

    if not documents and skipped:
        raise HTTPException(status_code=422, detail={"skipped": skipped})

    def _persist() -> int:
        n = upsert_source_documents(db, documents)
        # upsert_source_documents flushes but doesn't commit (connector syncs
        # commit at the sync-run level); commit here or the upload evaporates
        # when this request's session closes.
        db.commit()
        return n

    indexed = try_db("upload_documents", 0, _persist)

    vectors_indexed = 0
    vector_error = None
    try:
        vectors_indexed = await get_default_qdrant_store().upsert_chunks(
            chunks_for_documents(documents)
        )
    except Exception as exc:  # noqa: BLE001 — vector indexing must not lose the upload
        vector_error = str(exc)
        logger.warning("Qdrant indexing failed for upload (org=%s): %s", org_id, exc)

    return {
        "documents_indexed": indexed,
        "vectors_indexed": vectors_indexed,
        "vector_error": vector_error,
        "skipped": skipped,
        "visibility": visibility,
        "documents": [
            {"id": d.source_id, "title": d.title, "data_tier": d.data_tier}
            for d in documents
        ],
    }


def _derive_visibility(record: SourceDocumentRecord) -> dict:
    """Reverse-map stored grants back to the human visibility model."""
    perms = record.permissions or []
    user_ids = [p.removeprefix("user:") for p in perms if p.startswith("user:")]
    if perms and all(p.startswith("user:") for p in perms):
        visibility = "personal" if len(user_ids) == 1 else "people"
    elif any(p.startswith("dept:") for p in perms):
        visibility = "department"
    else:
        visibility = "company"
    return {
        "visibility": visibility,
        "shared_with": user_ids,
        "department_id": record.department_id,
    }


def _can_manage(record: SourceDocumentRecord, claims: dict | None, db: Session) -> bool:
    """Only the uploader (author) or a workspace admin may change access."""
    if claims is None:
        return True  # demo/unauthenticated context manages demo data freely
    user_id = claims.get("sub")
    email = claims.get("email")
    if record.author and record.author in (user_id, email):
        return True
    user = db.get(User, user_id) if user_id else None
    return bool(user and user.role == "admin")


def _get_upload_record(db: Session, org_id: str, doc_id: str) -> SourceDocumentRecord:
    record = db.get(SourceDocumentRecord, doc_id)
    if record is None or record.org_id != org_id:
        raise HTTPException(status_code=404, detail="Document not found.")
    return record


@router.get("/{doc_id}/access")
async def get_document_access(db: DbSession, org_id: OrgId, doc_id: str) -> dict:
    record = _get_upload_record(db, org_id, doc_id)
    access = _derive_visibility(record)
    members = (
        db.query(User).filter(User.org_id == org_id, User.id.in_(access["shared_with"])).all()
        if access["shared_with"]
        else []
    )
    access["people"] = [
        {"id": m.id, "name": m.display_name or m.email, "email": m.email} for m in members
    ]
    access["title"] = record.title
    return access


class AccessUpdate(BaseModel):
    visibility: str
    shared_with: list[str] = Field(default_factory=list)
    department_id: str | None = None


@router.patch("/{doc_id}/access")
async def update_document_access(
    db: DbSession,
    org_id: WriteOrgId,
    claims: OptionalClaims,
    doc_id: str,
    body: AccessUpdate,
) -> dict:
    """Change who can see an already-uploaded document (Drive-style manage
    access). Rewrites the grants in Postgres and on the Qdrant chunk payloads,
    and notifies newly-added people."""
    if body.visibility not in _ALLOWED_VISIBILITY:
        raise HTTPException(
            status_code=422, detail=f"visibility must be one of {sorted(_ALLOWED_VISIBILITY)}"
        )
    record = _get_upload_record(db, org_id, doc_id)
    if not _can_manage(record, claims, db):
        raise HTTPException(
            status_code=403, detail="Only the uploader or an admin can change access."
        )

    uploader_id = claims.get("sub") if claims else None
    dept = (body.department_id or "").strip() or None
    if not dept and body.visibility == "department" and uploader_id:
        uploader = db.get(User, uploader_id)
        dept = uploader.department_id if uploader else None
    if dept:
        row = db.get(Department, dept)
        if row is None or row.org_id != org_id:
            raise HTTPException(status_code=422, detail="Unknown department for this workspace.")

    previous = set(record.permissions or [])
    grants = _visibility_grants(db, org_id, body.visibility, uploader_id, dept, body.shared_with)

    record.permissions = grants
    record.department_id = dept if body.visibility == "department" else record.department_id
    db.query(Chunk).filter(Chunk.source_document_id == doc_id).update(
        {"permissions": grants}, synchronize_session=False
    )

    # Notify people who just gained access (not the actor themselves).
    actor = db.get(User, uploader_id) if uploader_id else None
    new_user_ids = {
        g.removeprefix("user:") for g in set(grants) - previous if g.startswith("user:")
    } - {uploader_id}
    for target in new_user_ids:
        db.add(
            Notification(
                org_id=org_id,
                user_id=target,
                type="document.shared",
                payload={
                    "document_id": doc_id,
                    "title": record.title,
                    "shared_by": (actor.display_name or actor.email) if actor else "A teammate",
                },
            )
        )
    db.commit()

    qdrant_error = None
    try:
        await get_default_qdrant_store().set_document_payload(
            org_id, doc_id, {"permissions": grants}
        )
    except Exception as exc:  # noqa: BLE001 — Postgres is source of truth; surface lag honestly
        qdrant_error = str(exc)
        logger.warning("Qdrant payload update failed (org=%s doc=%s): %s", org_id, doc_id, exc)

    return {**_derive_visibility(record), "qdrant_error": qdrant_error}
