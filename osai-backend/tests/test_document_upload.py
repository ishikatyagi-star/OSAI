"""Direct file upload into the knowledge base (POST /documents/upload)."""

from __future__ import annotations

import asyncio
import io
import json
import zipfile
from unittest.mock import AsyncMock, patch

import anyio
import pytest
import starlette.formparsers as starlette_formparsers
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from pypdf import PdfWriter

import api.ratelimit as limiter
from api import document_extraction
from api.main import app
from api.routes import documents as document_routes
from config import settings

client = TestClient(app)


def _upload(files, data=None):
    with patch("api.routes.documents.get_default_qdrant_store") as mock_store:
        mock_store.return_value.upsert_chunks = AsyncMock(return_value=len(files))
        return client.post("/documents/upload", files=files, data=data or {})


async def _raw_upload(
    body: bytes,
    content_type: str,
    *,
    receive_started: asyncio.Event | None = None,
) -> tuple[int, dict, int]:
    """Send exact multipart wire bytes and report application receive() calls."""
    messages: list[dict] = []
    receive_calls = 0
    body_sent = False

    async def receive() -> dict:
        nonlocal body_sent, receive_calls
        receive_calls += 1
        if receive_started is not None:
            receive_started.set()
        if body_sent:
            return {"type": "http.disconnect"}
        body_sent = True
        return {"type": "http.request", "body": body, "more_body": False}

    async def send(message: dict) -> None:
        messages.append(message)

    path = "/documents/upload"
    await app(
        {
            "type": "http",
            "asgi": {"version": "3.0", "spec_version": "2.3"},
            "http_version": "1.1",
            "method": "POST",
            "scheme": "http",
            "path": path,
            "raw_path": path.encode(),
            "root_path": "",
            "query_string": b"",
            "headers": [
                (b"host", b"testserver"),
                (b"content-type", content_type.encode("latin-1")),
                (b"content-length", str(len(body)).encode()),
            ],
            "client": ("198.51.100.10", 12345),
            "server": ("testserver", 80),
            "state": {},
        },
        receive,
        send,
    )
    status_code = next(
        message["status"] for message in messages if message["type"] == "http.response.start"
    )
    response_body = b"".join(
        message.get("body", b"") for message in messages if message["type"] == "http.response.body"
    )
    return status_code, json.loads(response_body), receive_calls


def test_upload_txt_ingests_document():
    resp = _upload(
        [("files", ("notes.txt", b"Quarterly planning notes: ship the pilot.", "text/plain"))]
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["documents_indexed"] == 1
    assert body["skipped"] == []
    doc = body["documents"][0]
    assert doc["title"] == "notes.txt"
    assert doc["data_tier"] == "normal"
    assert doc["id"].startswith("upload-")


def test_upload_docx_extracts_paragraphs():
    buf = io.BytesIO()
    d = DocxDocument()
    d.add_paragraph("Decision: adopt Qdrant for vectors.")
    d.save(buf)
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    resp = _upload([("files", ("decision.docx", buf.getvalue(), docx_mime))])
    assert resp.status_code == 200
    assert resp.json()["documents_indexed"] == 1


def test_upload_respects_tier():
    resp = _upload(
        [("files", ("salary.txt", b"Compensation bands for 2026.", "text/plain"))],
        data={"data_tier": "red"},
    )
    assert resp.status_code == 200
    assert resp.json()["documents"][0]["data_tier"] == "red"


def test_upload_default_visibility_is_personal():
    resp = _upload([("files", ("mine.txt", b"my notes", "text/plain"))])
    assert resp.status_code == 200
    assert resp.json()["visibility"] == "personal"


def test_upload_rejects_bad_visibility():
    resp = _upload(
        [("files", ("a.txt", b"x", "text/plain"))],
        data={"visibility": "everyone"},
    )
    assert resp.status_code == 422


def test_upload_company_visibility():
    resp = _upload(
        [("files", ("handbook.md", b"# Handbook", "text/markdown"))],
        data={"visibility": "company"},
    )
    assert resp.status_code == 200
    assert resp.json()["visibility"] == "company"


def test_upload_people_visibility_requires_recipients():
    resp = _upload(
        [("files", ("a.txt", b"x", "text/plain"))],
        data={"visibility": "people"},
    )
    assert resp.status_code == 422


def test_upload_department_visibility_requires_department():
    # Unauthenticated + no department_id: nothing to scope to.
    resp = _upload(
        [("files", ("a.txt", b"x", "text/plain"))],
        data={"visibility": "department"},
    )
    assert resp.status_code == 422


def test_upload_rejects_bad_tier():
    resp = _upload(
        [("files", ("a.txt", b"x", "text/plain"))],
        data={"data_tier": "topsecret"},
    )
    assert resp.status_code == 422


def test_unsupported_type_is_skipped_with_reason():
    resp = _upload([("files", ("binary.exe", b"\x00\x01", "application/octet-stream"))])
    assert resp.status_code == 422  # nothing ingestible
    detail = resp.json()["detail"]
    assert detail["skipped"][0]["filename"] == "binary.exe"


def test_mixed_batch_ingests_good_and_reports_skipped():
    resp = _upload(
        [
            ("files", ("ok.md", b"# Runbook\nRestart the worker.", "text/markdown")),
            ("files", ("bad.exe", b"\x00", "application/octet-stream")),
        ]
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["documents_indexed"] == 1
    assert len(body["skipped"]) == 1


def test_upload_rejects_too_many_files():
    resp = _upload([("files", (f"note-{index}.txt", b"x", "text/plain")) for index in range(11)])
    assert resp.status_code == 413


def test_upload_rejects_excessive_extracted_text():
    resp = _upload(
        [
            (
                "files",
                (
                    "oversized.txt",
                    b"x" * (document_extraction.MAX_EXTRACTED_TEXT_BYTES + 1),
                    "text/plain",
                ),
            )
        ]
    )
    assert resp.status_code == 422
    assert "Extracted text exceeds" in resp.json()["detail"]["skipped"][0]["reason"]


def test_upload_rejects_utf8_replacement_expansion_over_text_limit():
    # Each invalid byte decodes to U+FFFD, which is three bytes in UTF-8. The
    # wire payload is below 2 MiB while the extracted text is just over it.
    payload = b"\xff" * (document_extraction.MAX_EXTRACTED_TEXT_BYTES // 3 + 1)
    resp = _upload([("files", ("invalid-utf8.txt", payload, "text/plain"))])

    assert resp.status_code == 422
    assert "Extracted text exceeds" in resp.json()["detail"]["skipped"][0]["reason"]


def test_upload_rejects_input_over_per_file_limit():
    resp = _upload(
        [
            (
                "files",
                (
                    "large.txt",
                    b"x" * (document_extraction.MAX_FILE_BYTES + 1),
                    "text/plain",
                ),
            )
        ]
    )
    assert resp.status_code == 422
    assert resp.json()["detail"]["skipped"][0]["reason"] == "File exceeds 15 MB limit"


def test_upload_rejects_batch_over_total_input_limit(monkeypatch):
    monkeypatch.setattr(document_routes, "_MAX_BATCH_BYTES", 5)
    resp = _upload(
        [
            ("files", ("one.txt", b"123", "text/plain")),
            ("files", ("two.txt", b"456", "text/plain")),
        ]
    )
    assert resp.status_code == 413
    assert resp.json()["detail"] == "Upload exceeds 30 MB batch limit"


def test_upload_stream_is_bounded_before_form_materialization(monkeypatch):
    monkeypatch.setattr(document_routes, "_MAX_MULTIPART_BYTES", 512)
    resp = _upload([("files", ("large.txt", b"x" * 1024, "text/plain"))])
    assert resp.status_code == 413
    assert resp.json()["detail"] == "Upload exceeds 31 MB request limit"


def test_upload_rejects_pdf_over_page_limit():
    writer = PdfWriter()
    for _ in range(document_extraction.MAX_PDF_PAGES + 1):
        writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)

    resp = _upload([("files", ("long.pdf", buf.getvalue(), "application/pdf"))])
    assert resp.status_code == 422
    assert "PDF exceeds 250 page limit" in resp.json()["detail"]["skipped"][0]["reason"]


def test_upload_rejects_docx_archive_member_amplification():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as archive:
        for index in range(document_extraction.MAX_DOCX_MEMBERS + 1):
            archive.writestr(f"member-{index}.xml", b"x")

    resp = _upload(
        [
            (
                "files",
                (
                    "many-members.docx",
                    buf.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]
    )
    assert resp.status_code == 422
    assert "archive member limit" in resp.json()["detail"]["skipped"][0]["reason"]


def test_upload_rejects_docx_expansion_amplification():
    buf = io.BytesIO()
    block = b"x" * (64 * 1024)
    remaining = document_extraction.MAX_DOCX_EXPANDED_BYTES + 1
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as archive:
        with archive.open("word/document.xml", "w", force_zip64=True) as member:
            while remaining:
                chunk = block[:remaining]
                member.write(chunk)
                remaining -= len(chunk)

    resp = _upload(
        [
            (
                "files",
                (
                    "expanded.docx",
                    buf.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]
    )
    assert resp.status_code == 422
    assert "expanded content exceeds" in resp.json()["detail"]["skipped"][0]["reason"]


def test_upload_rejects_docx_extracted_text_amplification():
    buf = io.BytesIO()
    document = DocxDocument()
    document.add_paragraph("x" * (document_extraction.MAX_EXTRACTED_TEXT_BYTES + 1))
    document.save(buf)

    resp = _upload(
        [
            (
                "files",
                (
                    "long.docx",
                    buf.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]
    )
    assert resp.status_code == 422
    assert "Extracted text exceeds" in resp.json()["detail"]["skipped"][0]["reason"]


def test_upload_maps_parser_timeout_and_requests_worker_cancellation(monkeypatch):
    async def slow_parser(*_args, **_kwargs):
        assert document_routes._UPLOAD_ADMISSION_LIMITER.borrowed_tokens == 1
        await anyio.sleep(1)

    monkeypatch.setattr(document_routes, "_PARSER_TIMEOUT_SECONDS", 0.01)
    with patch(
        "api.routes.documents.anyio.to_process.run_sync", side_effect=slow_parser
    ) as run_sync:
        resp = _upload([("files", ("slow.txt", b"text", "text/plain"))])
    assert resp.status_code == 422
    assert "parsing exceeded" in resp.json()["detail"]["skipped"][0]["reason"]
    assert run_sync.await_args.kwargs["cancellable"] is True


@pytest.mark.asyncio
async def test_upload_rejects_multipart_prefix_impostor_without_consuming_body():
    status_code, payload, receive_calls = await _raw_upload(
        b"body must not be consumed",
        "multipart/form-datax; boundary=x",
    )

    assert status_code == 415
    assert payload == {"detail": "Content-Type must be multipart/form-data"}
    assert receive_calls == 0


@pytest.mark.asyncio
async def test_upload_rate_limit_rejects_without_consuming_body(monkeypatch):
    monkeypatch.setattr(settings, "env", "local")
    monkeypatch.setattr(limiter, "_memory_allow", lambda *_args, **_kwargs: False)

    status_code, payload, receive_calls = await _raw_upload(
        b"body must not be consumed",
        "multipart/form-data; boundary=x",
    )

    assert status_code == 429
    assert payload == {"detail": limiter._LIMIT_DETAIL}
    assert receive_calls == 0


@pytest.mark.asyncio
async def test_upload_admission_queues_before_second_body_is_consumed(monkeypatch):
    admission = anyio.CapacityLimiter(1)
    second_at_gate = asyncio.Event()

    class ObservedAdmission:
        attempts = 0

        async def __aenter__(self):
            self.attempts += 1
            if self.attempts == 2:
                second_at_gate.set()
            await admission.acquire()

        async def __aexit__(self, *_exc_info):
            admission.release()

    monkeypatch.setattr(document_routes, "_UPLOAD_ADMISSION_LIMITER", ObservedAdmission())
    first_parse_started = asyncio.Event()
    release_first = asyncio.Event()
    first_receive = asyncio.Event()
    second_receive = asyncio.Event()
    parse_calls = 0

    async def controlled_parse(parser):
        nonlocal parse_calls
        parse_calls += 1
        async for _chunk in parser.stream:
            break
        if parse_calls == 1:
            first_parse_started.set()
            await release_first.wait()
        raise document_routes.MultipartParseError("malformed")

    monkeypatch.setattr(document_routes.MultiPartParser, "parse", controlled_parse)
    first_task = asyncio.create_task(
        _raw_upload(
            b"first body",
            "multipart/form-data; boundary=x",
            receive_started=first_receive,
        )
    )
    await first_parse_started.wait()
    second_task = asyncio.create_task(
        _raw_upload(
            b"second body",
            "multipart/form-data; boundary=x",
            receive_started=second_receive,
        )
    )
    await second_at_gate.wait()
    try:
        assert first_receive.is_set()
        assert admission.borrowed_tokens == 1
        assert admission.statistics().tasks_waiting == 1
        assert not second_receive.is_set()
    finally:
        release_first.set()

    first_result, second_result = await asyncio.gather(first_task, second_task)
    assert first_result[0] == second_result[0] == 400
    assert first_result[2] == second_result[2] == 1


@pytest.mark.asyncio
async def test_upload_malformed_wire_body_returns_400():
    status_code, payload, receive_calls = await _raw_upload(
        b"not-a-multipart-boundary",
        "multipart/form-data; boundary=x",
    )

    assert status_code == 400
    assert payload == {"detail": "Malformed multipart body"}
    assert receive_calls == 1


@pytest.mark.asyncio
async def test_upload_truncated_wire_body_returns_400_and_closes_spool(monkeypatch):
    spools = []
    original_spooled_file = starlette_formparsers.SpooledTemporaryFile

    def tracked_spooled_file(*args, **kwargs):
        spool = original_spooled_file(*args, **kwargs)
        spools.append(spool)
        return spool

    monkeypatch.setattr(
        starlette_formparsers,
        "SpooledTemporaryFile",
        tracked_spooled_file,
    )
    truncated = (
        b'--x\r\nContent-Disposition: form-data; name="files"; filename="a.txt"\r\n'
        b"Content-Type: text/plain\r\n\r\nhello without a terminal boundary"
    )
    status_code, payload, receive_calls = await _raw_upload(
        truncated,
        "multipart/form-data; boundary=x",
    )

    assert status_code == 400
    assert payload == {"detail": "Malformed multipart body"}
    assert receive_calls == 1
    assert spools and all(spool.closed for spool in spools)


@pytest.mark.asyncio
async def test_upload_parse_error_closes_completed_file_spool(monkeypatch):
    spools = []
    original_spooled_file = starlette_formparsers.SpooledTemporaryFile

    def tracked_spooled_file(*args, **kwargs):
        spool = original_spooled_file(*args, **kwargs)
        spools.append(spool)
        return spool

    monkeypatch.setattr(
        starlette_formparsers,
        "SpooledTemporaryFile",
        tracked_spooled_file,
    )
    malformed_second_part = (
        b'--x\r\nContent-Disposition: form-data; name="files"; filename="a.txt"\r\n'
        b"Content-Type: text/plain\r\n\r\nfirst file\r\n"
        b"--x\r\ninvalid header\r\n\r\nsecond part\r\n--x--\r\n"
    )
    status_code, payload, receive_calls = await _raw_upload(
        malformed_second_part,
        "multipart/form-data; boundary=x",
    )

    assert status_code == 400
    assert payload == {"detail": "Malformed multipart body"}
    assert receive_calls == 1
    assert spools and all(spool.closed for spool in spools)


@pytest.mark.asyncio
async def test_upload_requires_auth_without_consuming_body():
    from db.session import get_org_id, require_writable_org

    app.dependency_overrides.pop(get_org_id, None)
    app.dependency_overrides.pop(require_writable_org, None)

    status_code, payload, receive_calls = await _raw_upload(
        b"body must not be consumed",
        "multipart/form-data; boundary=x",
    )

    assert status_code == 401
    assert payload == {"detail": "Authentication required."}
    assert receive_calls == 0


def _first_doc_id(resp) -> str:
    return resp.json()["documents"][0]["id"]


def test_access_roster_requires_authenticated_manager():
    up = _upload([("files", ("plan.txt", b"Q3 plan", "text/plain"))])
    doc_id = _first_doc_id(up)
    # The test auth override has no user claims. Even for a company-visible file,
    # the full recipient/email roster remains uploader-or-admin metadata.
    resp = client.get(f"/documents/{doc_id}/access")
    assert resp.status_code == 403


def test_access_patch_updates_visibility():
    up = _upload([("files", ("plan2.txt", b"Q4 plan", "text/plain"))])
    doc_id = _first_doc_id(up)
    with patch("api.routes.documents.get_default_qdrant_store") as mock_store:
        mock_store.return_value.set_document_payload = AsyncMock(return_value=None)
        resp = client.patch(f"/documents/{doc_id}/access", json={"visibility": "company"})
    assert resp.status_code == 200
    assert resp.json()["visibility"] == "company"
    assert resp.json()["qdrant_error"] is None


def test_access_patch_rejects_bad_visibility_and_unknown_doc():
    up = _upload([("files", ("plan3.txt", b"notes", "text/plain"))])
    doc_id = _first_doc_id(up)
    assert (
        client.patch(f"/documents/{doc_id}/access", json={"visibility": "everyone"}).status_code
        == 422
    )
    assert (
        client.patch("/documents/upload-nope/access", json={"visibility": "company"}).status_code
        == 404
    )


def test_notifications_empty_for_unauthenticated():
    assert client.get("/notifications").status_code == 200
    assert client.get("/notifications").json() == []
