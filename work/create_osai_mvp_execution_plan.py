from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Admin\Documents\Codex\2026-06-02\files-mentioned-by-the-user-osai\outputs")
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "OSAI_MVP_Execution_Plan_Updated_Stack.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr[idx], "1F4E79")
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(f"\n{body}")
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 24, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 12.5, "2F5597"),
        ("Heading 3", 11.2, "404040"),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("OSAI MVP Execution Plan - Updated Stack - June 2026")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def build():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("OSAI MVP Execution Plan\n")
    sub = title.add_run("Updated Stack, Connector Strategy, and Pilot Build Roadmap")
    sub.font.size = Pt(13)
    sub.font.bold = False
    sub.font.color.rgb = RGBColor(89, 89, 89)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for the OSAI founding team and build agent | Date: June 2, 2026")

    add_callout(
        doc,
        "Build thesis",
        "Ship a connector-first OSAI MVP quickly by reusing proven open-source architecture patterns, "
        "using Composio for OAuth/tool execution where possible, and building only the opinionated OSAI layer: "
        "source ingestion, permission-aware memory, context retrieval, workflow execution, and the internal dashboard.",
    )

    doc.add_heading("1. Executive Decision", level=1)
    doc.add_paragraph(
        "The original brief is directionally right, but Section 12 is too narrow for the actual pilot. "
        "It treats the MVP as a Zoom-to-Linear workflow. The updated MVP should be a connector-first operating layer "
        "with one polished workflow and a small set of high-value knowledge and business-system connectors."
    )
    add_bullets(
        doc,
        [
            "Do not build a universal connector marketplace in the first pilot.",
            "Do build a stable connector registry that lets us add sources without rewriting workflows.",
            "Use open-source repos as reference or internal scaffolding, but avoid GPL contamination in distributed SaaS code.",
            "Prioritize integrations that prove the pain: scattered communication, scattered knowledge, unclear ownership, and manual internal ops.",
            "Pilot success equals trusted context retrieval plus visible execution: tasks created, support tickets summarized, HR/finance answers surfaced, and dashboard clarity.",
        ],
    )

    doc.add_heading("2. MVP Product Surface", level=1)
    add_table(
        doc,
        ["Surface", "What ships in MVP", "What waits"],
        [
            [
                "Internal dashboard",
                "One admin dashboard showing connected tools, recent syncs, workflow runs, extracted action items, and source health.",
                "Custom dashboard builder, advanced analytics, and multi-role permissions.",
            ],
            [
                "Org brain search",
                "Ask questions across Slack/email/docs/Notion/Freshdesk with citations and source links.",
                "Full enterprise graph and deep personal graph.",
            ],
            [
                "Action workflow",
                "Meeting/action-item extraction that creates Notion tasks or Linear/GitHub tickets and posts a Slack/email summary.",
                "General workflow builder and complex approval routing.",
            ],
            [
                "Tool setup",
                "Guided connect screens for P0 connectors, using Composio where possible and direct APIs where necessary.",
                "Marketplace, billing gates, public app directory.",
            ],
            [
                "Governance",
                "Simple org-level data tier flags: Normal, Amber, Red. Red data routes to local model when configured.",
                "Full RBAC, SAML/SSO, audit exports.",
            ],
        ],
        widths=[1.5, 3.2, 2.4],
    )

    doc.add_heading("3. Integration Priority", level=1)
    add_table(
        doc,
        ["Priority", "Connector", "Reason", "Implementation path"],
        [
            ["P0", "Slack", "Primary scattered-communication source and notification destination.", "Composio or Slack API Events/OAuth."],
            ["P0", "Email: Gmail and Outlook", "Decisions and client context live in inboxes.", "Google APIs and Microsoft Graph; use Composio if pilot auth is faster."],
            ["P0", "Google Drive", "Core knowledge base for docs, sheets, decks, PDFs.", "Google Drive API; sync metadata and extract file text."],
            ["P0", "Microsoft OneDrive/SharePoint", "Office files and enterprise documents.", "Microsoft Graph files APIs."],
            ["P0", "Notion", "Task management and internal docs.", "Official Notion API; create tasks/pages and ingest pages/databases."],
            ["P0", "Freshdesk", "Customer support context and ticket trends.", "Freshdesk REST API v2."],
            ["P1", "Keka", "HR answers: leave, attendance, employee profile, policies.", "Official Keka API; read-only first."],
            ["P1", "TallyPrime", "Finance answers and lightweight reporting.", "TallyPrime XML/HTTP or JSON integration depending on installed version."],
            ["P1", "WhatsApp Business", "High-signal informal customer/team communication.", "Meta WhatsApp Cloud API; likely read/summarize first, write later."],
            ["P1", "Cursor / code context", "Tech team context and code-agent workflows.", "Expose OSAI memory as MCP; optionally ingest repo docs/issues."],
            ["P2", "ChatGPT / Claude memory bridge", "Knowledge currently created inside AI tools.", "MCP bridge, exports, or user-shared artifacts; do not depend on unavailable private data APIs."],
        ],
        widths=[0.7, 1.45, 2.55, 2.4],
    )

    doc.add_heading("4. Updated Technical Stack", level=1)
    add_table(
        doc,
        ["Layer", "Decision", "Configuration"],
        [
            ["Backend", "Python 3.12 or 3.13, FastAPI, Pydantic v2", "Use uv for dependency locking. Keep async FastAPI routes and sync Celery workers."],
            ["Workers", "Celery + Redis", "Queues: ingest, transcribe, extract, execute, maintenance. Use retries, idempotency keys, and dead-letter handling."],
            ["Frontend", "Next.js current stable on Node 24 LTS", "App Router, server actions where useful, Tailwind/shadcn only if already accepted by repo."],
            ["Database", "PostgreSQL via Supabase for pilot", "Tables: orgs, users, connectors, sync_runs, documents, chunks, workflow_runs, action_items, audit_events."],
            ["Vector store", "Qdrant", "One collection per org or tenant-prefixed collection names. Payload filters enforce org_id, source_id, data_tier, permissions."],
            ["Object storage", "Supabase Storage or S3-compatible storage", "Store raw transcripts, extracted text artifacts, file snapshots, and replay payloads."],
            ["Connector execution", "Composio first, direct API second", "Use Composio for OAuth/actions where supported. Direct APIs for Keka, TallyPrime, and any unsupported pilot connector."],
            ["LLM gateway", "LiteLLM or provider adapter behind internal model router", "Default: GPT-5-class or Claude Sonnet-class extraction model. Fallback: cheaper mini/Haiku-class model. Red data: Ollama or approved local model."],
            ["Transcription", "OpenAI audio transcription with whisper-1 initially", "Use verbose_json and word timestamps when needed. Self-host Whisper only when volume justifies ops cost."],
            ["Observability", "Structured logs + Sentry/OpenTelemetry", "Every workflow_run has trace_id, connector_run_id, source refs, model, token usage, and execution outcome."],
            ["Deployment", "Docker Compose for pilot; Railway/Fly/Render or single VM acceptable", "Do not overbuild Kubernetes for pilot. Keep compose parity with production env vars."],
        ],
        widths=[1.25, 2.35, 3.45],
    )

    doc.add_heading("5. Open-Source Reuse Strategy", level=1)
    add_table(
        doc,
        ["Repo / project", "Use level", "Why", "Guardrail"],
        [
            ["onyx-dot-app/onyx", "Reference and selective pattern reuse", "Best reference for enterprise connectors, permission-aware RAG, and chunk metadata.", "Do not copy UI wholesale; use connector and permission ideas."],
            ["ComposioHQ/composio", "Direct dependency", "Fastest path to OAuth and tool actions across many apps.", "Pin SDK versions and test auth flows per connector."],
            ["qdrant/qdrant", "Direct dependency", "Production-ready vector DB with filters and hybrid search support.", "Keep schema simple until pilot usage proves graph needs."],
            ["openai/whisper", "Fallback/self-host later", "Open-source transcription path if API costs rise.", "Use managed API first unless volume is high."],
            ["BerriAI/litellm", "Direct or adapter reference", "Unified model routing and fallbacks.", "Pin known-good versions; monitor security advisories."],
            ["tinyhumansai/openhuman", "Reference only", "Useful memory tree, chunking, and personal AI patterns.", "GPL-3.0: do not directly fork into closed/distributed SaaS code."],
            ["agiresearch/AIOS", "Reference only", "Future agent scheduling and memory isolation ideas.", "Not MVP-critical."],
            ["celery/celery", "Direct dependency", "Stable async job orchestration for sync/extract/execute chains.", "Prefer simple queues over complex canvas until needed."],
        ],
        widths=[1.75, 1.4, 2.65, 2.0],
    )

    doc.add_heading("6. Target Architecture", level=1)
    add_numbered(
        doc,
        [
            "Connector registry stores metadata for each integration: capabilities, auth mode, scopes, sync cadence, supported actions, and data tier.",
            "Ingestion workers pull or receive webhook events, normalize source records, and write raw source metadata plus extracted text.",
            "Chunker creates 500-1200 token chunks with source_id, doc_id, source_type, org_id, user_id, data_tier, permissions, created_at, updated_at, and content_preview.",
            "Embedding worker writes chunks to Qdrant and keeps Postgres as the source of truth for provenance and permissions.",
            "Retriever performs hybrid search with org_id and permission filters before any model call.",
            "Workflow engine takes a trigger, fetches context, calls the model router, validates JSON with Pydantic, and executes connector actions.",
            "Dashboard reads workflow_runs, connector health, action_items, and source coverage; it never calls third-party APIs directly from the browser.",
        ],
    )

    doc.add_heading("7. Connector Interface Contract", level=1)
    doc.add_paragraph(
        "All connectors should implement a common interface so the pilot does not become a pile of one-off integrations."
    )
    add_table(
        doc,
        ["Method", "Purpose", "Required for MVP"],
        [
            ["auth_status(org_id)", "Return connected/disconnected/error plus scopes.", "Yes"],
            ["sync(org_id, cursor)", "Fetch changed records and return normalized SourceDocument objects.", "Yes for read connectors"],
            ["get_permissions(document)", "Map source permissions into OSAI payload filters.", "Yes for docs/email/Notion"],
            ["search(query)", "Optional source-native lookup when vector context is insufficient.", "Optional"],
            ["execute_action(action)", "Create task, ticket, comment, message, or update.", "Yes for Notion/Freshdesk/Slack"],
            ["healthcheck()", "Validate credentials, rate limits, and API availability.", "Yes"],
        ],
        widths=[1.7, 3.7, 1.5],
    )

    doc.add_heading("8. MVP Build Phases", level=1)
    add_table(
        doc,
        ["Phase", "Duration", "Build output", "Exit criteria"],
        [
            ["0. Repo setup", "Day 1", "Monorepo or two repos, Docker Compose, env templates, CI lint/test.", "FastAPI, Next.js, Postgres, Redis, Qdrant all boot locally."],
            ["1. Core data model", "Day 2", "Postgres schema, connector registry, workflow_run schema, audit events.", "Seed org and first admin can exist; migrations run cleanly."],
            ["2. P0 connector auth", "Days 3-4", "Slack, Google, Microsoft, Notion, Freshdesk connection records.", "At least 3 P0 connectors show connected status in dashboard."],
            ["3. Ingestion pipeline", "Days 5-7", "Sync workers, chunker, embeddings, Qdrant writes.", "Search returns cited results across at least Slack/docs/Notion or Freshdesk."],
            ["4. Workflow engine", "Days 8-10", "Meeting/action-item extractor and task/ticket execution.", "A transcript produces validated action items and creates Notion tasks or tickets."],
            ["5. Dashboard", "Days 11-12", "Connector health, sync log, search, workflow log, action detail pages.", "Pilot user can understand what happened without reading logs."],
            ["6. Hardening", "Days 13-14", "Retries, DLQ, permission filters, model fallback, error UI, seed demo data.", "Pilot demo is repeatable end to end with real or staged data."],
        ],
        widths=[1.2, 0.85, 3.15, 2.4],
    )

    doc.add_heading("9. Immediate Sprint Backlog", level=1)
    add_table(
        doc,
        ["Owner", "Task", "Definition of done"],
        [
            ["Build agent", "Create backend scaffold with FastAPI, uv, Pydantic v2, Alembic, Celery, Redis, Qdrant client.", "Tests pass; /health returns DB/Redis/Qdrant status."],
            ["Build agent", "Create connector base classes and registry.", "Mock connector syncs and writes normalized documents."],
            ["Build agent", "Implement Notion connector first.", "Can ingest selected database/pages and create task page/action item."],
            ["Build agent", "Implement Slack connector second.", "Can ingest channel messages and post workflow summary."],
            ["Build agent", "Implement Freshdesk connector third.", "Can ingest tickets and summarize customer-support context."],
            ["Build agent", "Implement search API.", "Query endpoint returns answer, citations, source records, and confidence."],
            ["Build agent", "Implement workflow runner.", "Transcript or pasted meeting notes creates validated action items and executes task creation."],
            ["Frontend", "Build dashboard skeleton.", "Connectors, sync runs, search, workflow runs, and action item detail pages are usable."],
            ["Founder/CTO", "Confirm pilot connectors and credentials.", "Access exists for at least 3 real data sources."],
        ],
        widths=[1.25, 3.3, 2.7],
    )

    doc.add_heading("10. Model and Prompt Configuration", level=1)
    add_bullets(
        doc,
        [
            "Do not hardcode a single model string inside business logic. Use an internal model_router config.",
            "Use a strong reasoning/extraction model for action-item extraction and policy-sensitive synthesis.",
            "Use a cheaper model for classification, source tagging, title generation, and summaries where failures are low-risk.",
            "Use local Ollama or approved local inference only for Red-tier data when the pilot requires it.",
            "Cache stable system prompts where provider support exists. Store prompt version, model, and schema version on every workflow_run.",
        ],
    )
    add_table(
        doc,
        ["Use case", "Primary", "Fallback", "Validation"],
        [
            ["Action extraction", "GPT-5-class or Claude Sonnet-class", "Alternate strong provider model", "Pydantic ActionItem schema; reject invalid JSON."],
            ["Search answer synthesis", "Strong general model", "Mini/Haiku-class model for short answers", "Must include citations; no citation means no confident answer."],
            ["Classification", "Mini/Haiku-class", "Local model if low-risk", "Enum-only outputs."],
            ["Red-tier data", "Local model", "No cloud fallback without explicit org setting", "Audit event required."],
        ],
        widths=[1.7, 2.0, 2.0, 1.8],
    )

    doc.add_heading("11. Pilot Acceptance Criteria", level=1)
    add_bullets(
        doc,
        [
            "Connect at least three real pilot systems: one communication source, one knowledge source, and one business system.",
            "Ask a natural-language question and get a cited answer from indexed pilot data.",
            "Run the action-item workflow and create a real Notion task, support ticket update, or engineering ticket.",
            "Dashboard shows connector status, last sync time, failures, workflow outputs, and source citations.",
            "Permission filter prevents cross-org or unauthorized source leakage in test cases.",
            "Every connector action has an audit event with actor, source, payload hash, status, and timestamp.",
            "Failure paths are visible: expired credentials, rate limits, model failure, invalid LLM JSON, and action execution failure.",
        ],
    )

    doc.add_heading("12. Risks and Mitigations", level=1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Connector sprawl", "MVP slips because every tool behaves differently.", "Use P0/P1 priority and a common connector interface. Build 3-5 connectors well."],
            ["Auth delays", "Pilot blocked by OAuth app setup.", "Use Composio where it saves time; direct API keys for internal pilot when safe."],
            ["Permission leakage", "Trust failure.", "Filter retrieval by org_id, source_id, permissions, and data_tier before model calls."],
            ["LLM hallucination", "Bad answers or wrong actions.", "Require citations for answers and schema validation plus human-visible action previews for risky writes."],
            ["GPL contamination", "Licensing issue.", "Do not fork GPL OpenHuman into proprietary distributed product; use as architecture reference."],
            ["Tally/Keka API variance", "Finance/HR connector takes longer than expected.", "Make them read-only P1. Start with API proof of connectivity and CSV/manual fallback."],
            ["Pilot data messiness", "Search feels low quality.", "Ship source health, recency metadata, and citation UI so users see provenance."],
        ],
        widths=[1.75, 2.15, 3.25],
    )

    doc.add_heading("13. Updated Section 12 Replacement Text", level=1)
    doc.add_paragraph(
        "Replace the old coding-agent instructions with this shorter directive:"
    )
    add_callout(
        doc,
        "Coding-agent directive",
        "Build OSAI as a connector-first operating layer. Implement a shared connector interface, a permission-aware memory pipeline, "
        "a model-router-backed workflow engine, and a dashboard that makes syncs and actions visible. Use Composio for OAuth and actions "
        "where it accelerates delivery. Use direct APIs for connectors Composio does not cover well, especially Keka, TallyPrime, and pilot-specific systems. "
        "The first workflow must turn meeting notes/transcripts into validated action items and create tasks/tickets in Notion, Linear, GitHub, or Freshdesk depending on pilot need.",
    )

    doc.add_heading("14. Source Notes Checked", level=1)
    add_bullets(
        doc,
        [
            "Composio docs and GitHub: current Python SDK and open-source toolkit direction.",
            "Qdrant docs/GitHub: vector store and Python client usage.",
            "Onyx GitHub/docs: connector and permission-aware enterprise RAG patterns.",
            "OpenAI docs: audio transcription with whisper-1 and current model/API guidance.",
            "Anthropic docs: Claude Sonnet-class model naming and API model direction.",
            "Slack, Google, Microsoft Graph, Notion, Freshdesk, Keka, Meta WhatsApp, and TallyPrime official docs for integration feasibility.",
            "LiteLLM docs/GitHub: routing/fallback model gateway, with version pinning due to security advisories.",
        ],
    )

    doc.add_heading("15. Canonical URLs", level=1)
    urls = [
        "https://github.com/onyx-dot-app/onyx",
        "https://docs.onyx.app/admin/connectors",
        "https://github.com/ComposioHQ/composio",
        "https://docs.composio.dev/",
        "https://github.com/qdrant/qdrant",
        "https://qdrant.tech/documentation/",
        "https://github.com/openai/whisper",
        "https://platform.openai.com/docs/guides/speech-to-text",
        "https://docs.litellm.ai/",
        "https://developers.notion.com/",
        "https://developers.freshdesk.com/api/",
        "https://apidocs.keka.com/",
        "https://help.tallysolutions.com/integration-with-tallyprime/",
        "https://developers.facebook.com/docs/whatsapp/cloud-api",
        "https://developers.google.com/drive/api",
        "https://learn.microsoft.com/graph/",
        "https://api.slack.com/",
    ]
    for url in urls:
        doc.add_paragraph(url, style="List Bullet")

    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
